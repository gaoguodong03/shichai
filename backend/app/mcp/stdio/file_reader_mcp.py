#!/usr/bin/env python3
"""文件读取与写入 MCP Server（本地 stdio，统一替代 Filesystem MCP）

提供 read_file、write_file、read_pdf、read_docx、read_xlsx、list_allowed_directories、list_directory。
路径相对 backend/data（如 data/users/<user>/agent-outputs/workspaces/...）。环境变量 FILE_READER_DATA_ROOT 可覆盖。
无需 Node.js，仅需 Python 与 pypdf/python-docx/openpyxl。
"""
import os
from pathlib import Path

from mcp.server.fastmcp import FastMCP

# 本文件位于 backend/app/mcp/stdio/，backend 根为向上 3 级
BACKEND_DIR = Path(__file__).resolve().parents[3]
DEFAULT_ROOT = BACKEND_DIR / "data"
ROOT = Path(os.getenv("FILE_READER_DATA_ROOT", str(DEFAULT_ROOT))).resolve()


def _resolve_path(relative_path: str) -> Path:
    """将相对路径解析为绝对路径，且必须落在 ROOT（默认 backend/data）内。"""
    raw = (relative_path or "").strip().strip("/").replace("..", "")
    if not raw:
        return None
    if raw.startswith("data/"):
        raw = raw[len("data/") :].lstrip("/")
    full = (ROOT / raw).resolve()
    if not str(full).startswith(str(ROOT)):
        return None
    return full


def _safe_read_text(path: Path, encoding: str = "utf-8") -> str:
    """读取纯文本文件"""
    try:
        return path.read_text(encoding=encoding)
    except UnicodeDecodeError:
        return f"错误：文件 {path.name} 不是 UTF-8 文本，请使用 read_pdf 或 read_docx 等工具。"
    except Exception as e:
        return f"错误：读取文件失败 - {e}"


mcp = FastMCP("File Reader Server")


@mcp.tool()
def write_file(path: str, content: str) -> str:
    """将文本内容写入文件。path 为相对路径（如 notes/report.md 或 workspace-write-test.txt）。

    当用户要求「保存到工作区」「写入文件」「把内容写到 xxx」时使用。path 与 content 必填；若未传 content 会报错。"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content or "", encoding="utf-8")
        rel = str(p.relative_to(ROOT)).replace("\\", "/")
        return f"已写入：{rel}"
    except Exception as e:
        return f"错误：写入文件失败 - {e}"


@mcp.tool()
def list_allowed_directories() -> str:
    """返回当前允许读写的根目录（默认 backend/data）。Agent 可在此目录及其子目录下使用 read_file、write_file 等。"""
    return f"Allowed directories:\n{ROOT}"


@mcp.tool()
def list_directory(path: str = "") -> str:
    """列出目录下的文件与子目录名。path 为空或 '.' 表示 data 根下目录；可为 users/<用户>/agent-outputs/workspaces/某会话ID 等相对路径。"""
    p = _resolve_path(path or ".")
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：路径不存在：{path}"
    if not p.is_dir():
        return f"错误：{path} 不是目录。"
    try:
        names = sorted(p.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
        lines = [f"{'[DIR]  ' if x.is_dir() else '       '}{x.name}" for x in names]
        return "\n".join(lines) if lines else "（空目录）"
    except Exception as e:
        return f"错误：列出目录失败 - {e}"


@mcp.tool()
def read_file(path: str) -> str:
    """读取纯文本文件（txt、md、json、yaml 等）。path 为相对路径（如 report.txt 或带子目录的 output/pages/xxx/text.md）。

    当本轮消息带有结构化附件时，path 使用附件中的工作区相对路径；若是 PDF、DOC、Excel，请使用 read_pdf、read_docx、read_xlsx。"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _safe_read_text(p)


@mcp.tool()
def read_pdf(path: str, max_pages: int = 50) -> str:
    """从 PDF 文件提取文本。path 为相对路径（如 report.pdf）。

    当结构化附件或用户明确路径指向 PDF 时，使用此工具。

    Args:
        path: 相对路径，如 report.pdf
        max_pages: 最多提取页数，默认 50，避免超出 token 限制"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _extract_pdf_text(p, max_pages)


def _extract_pdf_text(path: Path, max_pages: int) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "错误：未安装 pypdf。请运行: pip install pypdf"
    try:
        reader = PdfReader(path)
        total = len(reader.pages)
        if total == 0:
            return "（PDF 无页面内容）"
        texts = []
        for page in reader.pages[:max_pages]:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n\n".join(texts) + (f"\n\n（已截断，共 {total} 页，仅提取前 {max_pages} 页）" if total > max_pages else "")
    except Exception as e:
        return f"错误：读取 PDF 失败 - {e}"


@mcp.tool()
def read_docx(path: str) -> str:
    """从 DOCX 文件提取文本。path 为相对路径（如 report.docx）。

    当结构化附件或用户明确路径指向 DOCX 时，使用此工具。"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _extract_docx_text(p)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "错误：未安装 python-docx。请运行: pip install python-docx"
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts) if parts else "（DOCX 无文本内容）"
    except Exception as e:
        return f"错误：读取 DOCX 失败 - {e}"


@mcp.tool()
def read_xlsx(path: str, max_rows: int = 500) -> str:
    """从 Excel 文件（xlsx）提取文本。path 为相对路径（如 data.xlsx）。

    当结构化附件或用户明确路径指向 Excel 文件时，使用此工具。

    Args:
        path: 相对路径，如 data.xlsx
        max_rows: 每个 sheet 最多提取行数，默认 500"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _extract_xlsx_text(p, max_rows)


def _extract_xlsx_text(path: Path, max_rows: int) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "错误：未安装 openpyxl。请运行: pip install openpyxl"
    try:
        wb = load_workbook(path, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            parts.append(f"=== Sheet: {sheet_name} ===")
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    rows.append(f"（已截断，仅显示前 {max_rows} 行）")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
            parts.append("\n".join(rows))
        return "\n\n".join(parts) if parts else "（Excel 无内容）"
    except Exception as e:
        return f"错误：读取 Excel 失败 - {e}"


if __name__ == "__main__":
    mcp.run(transport="stdio")
