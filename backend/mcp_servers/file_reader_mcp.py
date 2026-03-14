#!/usr/bin/env python3
"""文件读取与写入 MCP Server（本地 stdio）

提供 read_file、read_pdf、read_docx、read_xlsx、write_file。
路径均相对 AGENT_OUTPUTS_DIR；在 DHA 中按会话使用时，后端会将 path 重写为 workspaces/{session_id}/ 下，实现工作区隔离。
"""
import os
from pathlib import Path

from mcp.server.fastmcp import FastMCP

# 与 backend/app/api/files 一致
SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent
DEFAULT_ROOT = BACKEND_DIR / "data" / "agent-outputs"
ROOT = Path(os.getenv("AGENT_OUTPUTS_DIR", str(DEFAULT_ROOT))).resolve()


def _resolve_path(relative_path: str) -> Path:
    """将相对路径解析为绝对路径，且必须落在 ROOT 内。接受相对 ROOT 的路径，或带 data/agent-outputs/ 前缀的路径（与 session wrapper 一致）。"""
    raw = (relative_path or "").strip().strip("/").replace("..", "")
    if not raw:
        return None
    # 后端 wrapper 可能传入 data/agent-outputs/workspaces/{session_id}/xxx，去掉与 ROOT 同名的前缀
    prefix = "data/agent-outputs/"
    if raw.startswith(prefix):
        raw = raw[len(prefix) :].lstrip("/")
    normalized = raw
    full = (ROOT / normalized).resolve()
    if not str(full).startswith(str(ROOT)):
        return None
    return full


def _safe_read_text(path: Path, encoding: str = "utf-8") -> str:
    """读取纯文本文件"""
    try:
        return path.read_text(encoding=encoding)
    except UnicodeDecodeError:
        return f"错误：文件 {path.name} 不是 UTF-8 文本，请使用 read_pdf 或 read_docx 等工具。"
    except Exception as e:
        return f"错误：读取文件失败 - {e}"


mcp = FastMCP("File Reader Server", description="在工作区内读取与写入文件（文本、PDF、DOCX、XLSX）")


@mcp.tool()
def write_file(path: str, content: str) -> str:
    """将文本内容写入文件。path 为相对路径（如 notes/report.md 或 workspace-write-test.txt）。

    当用户要求「保存到工作区」「写入文件」「把内容写到 xxx」时使用。path 与 content 必填；若未传 content 会报错。"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content or "", encoding="utf-8")
        rel = str(p.relative_to(ROOT)).replace("\\", "/")
        return f"已写入：{rel}"
    except Exception as e:
        return f"错误：写入文件失败 - {e}"


@mcp.tool()
def read_file(path: str) -> str:
    """读取纯文本文件（txt、md、json、yaml 等）。path 为相对路径（如 report.txt 或 workspaces/{会话ID}/notes.md）。

    当用户消息中出现【文件引用：path】时，若文件是文本格式，使用此工具；若是 PDF、DOC、Excel，请使用 read_pdf、read_docx、read_xlsx。"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _safe_read_text(p)


@mcp.tool()
def read_pdf(path: str, max_pages: int = 50) -> str:
    """从 PDF 文件提取文本。path 为相对路径（如 report.pdf）。

    当用户消息中出现【文件引用：xxx.pdf】时，使用此工具。

    Args:
        path: 相对路径，如 report.pdf
        max_pages: 最多提取页数，默认 50，避免超出 token 限制"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _extract_pdf_text(p, max_pages)


def _extract_pdf_text(path: Path, max_pages: int) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "错误：未安装 pypdf。请运行: pip install pypdf"
    try:
        reader = PdfReader(path)
        total = len(reader.pages)
        if total == 0:
            return "（PDF 无页面内容）"
        texts = []
        for page in reader.pages[:max_pages]:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n\n".join(texts) + (f"\n\n（已截断，共 {total} 页，仅提取前 {max_pages} 页）" if total > max_pages else "")
    except Exception as e:
        return f"错误：读取 PDF 失败 - {e}"


@mcp.tool()
def read_docx(path: str) -> str:
    """从 DOCX 文件提取文本。path 为相对路径（如 report.docx）。

    当用户消息中出现【文件引用：xxx.docx】时，使用此工具。"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _extract_docx_text(p)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "错误：未安装 python-docx。请运行: pip install python-docx"
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts) if parts else "（DOCX 无文本内容）"
    except Exception as e:
        return f"错误：读取 DOCX 失败 - {e}"


@mcp.tool()
def read_xlsx(path: str, max_rows: int = 500) -> str:
    """从 Excel 文件（xlsx）提取文本。path 为相对路径（如 data.xlsx）。

    当用户消息中出现【文件引用：xxx.xlsx】时，使用此工具。

    Args:
        path: 相对路径，如 data.xlsx
        max_rows: 每个 sheet 最多提取行数，默认 500"""
    p = _resolve_path(path)
    if not p:
        return "错误：无效路径或路径超出允许范围。"
    if not p.exists():
        return f"错误：文件不存在：{path}"
    if p.is_dir():
        return f"错误：{path} 是目录，无法读取。"
    return _extract_xlsx_text(p, max_rows)


def _extract_xlsx_text(path: Path, max_rows: int) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "错误：未安装 openpyxl。请运行: pip install openpyxl"
    try:
        wb = load_workbook(path, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            parts.append(f"=== Sheet: {sheet_name} ===")
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    rows.append(f"（已截断，仅显示前 {max_rows} 行）")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append("\t".join(cells))
            parts.append("\n".join(rows))
        return "\n\n".join(parts) if parts else "（Excel 无内容）"
    except Exception as e:
        return f"错误：读取 Excel 失败 - {e}"


if __name__ == "__main__":
    mcp.run(transport="stdio")
